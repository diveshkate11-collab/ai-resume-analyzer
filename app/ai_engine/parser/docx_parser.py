from docx import Document

class DOCXParser:
    """
    Extracts text from DOCX files.
    """

    @staticmethod
    def extract_text(file_path: str) -> dict:
        """
        Reads a DOCX file and extracts its text.

        Args:
            file_path (str): Path to the DOCX file.

        Returns:
            dict: Extracted text with metadata.
        """

        try:
            document = Document(file_path)
            extracted_text = "\n".join(
                paragraph.text
                for paragraph in document.paragraphs
            )

            return {
                "paragraphs": len(document.paragraphs),
                "characters": len(extracted_text),
                "text": extracted_text,
            }

        except Exception as e:
            raise Exception(f"Error reading DOCX: {e}")